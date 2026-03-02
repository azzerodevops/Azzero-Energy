"""Generate DOCX reports using python-docx.

Produces a professionally formatted energy-analysis report with branded
colours, styled tables, and embedded matplotlib charts.

The main entry point is ``generate_docx(data)`` which returns a ``Path``
to the saved ``.docx`` file inside a temporary directory.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from .models import ReportData
from .charts import generate_all_charts

# ---------------------------------------------------------------------------
# Brand colours for document elements
# ---------------------------------------------------------------------------

PRIMARY_BLUE = RGBColor(0x00, 0x97, 0xD7)
SECONDARY_GREEN = RGBColor(0x00, 0xB8, 0x94)
DARK_BG = RGBColor(0x12, 0x18, 0x27)
MUTED_TEXT = RGBColor(0x94, 0xA3, 0xB8)

# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading with branded blue colour."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = PRIMARY_BLUE


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a formatted table with header styling."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = cell_text
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def _fmt_eur(value: float) -> str:
    """Format a number as Euro with dot-separated thousands."""
    return f"\u20ac {value:,.0f}".replace(",", ".")


def _fmt_pct(value: float | None) -> str:
    """Format a ratio (0-1) as percentage, or em-dash for *None*."""
    if value is None:
        return "\u2014"
    return f"{value * 100:.1f}%"


def _add_centered_image(doc: Document, image_path: str, width: Inches) -> None:
    """Insert an image and centre the containing paragraph."""
    doc.add_picture(image_path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def generate_docx(data: ReportData) -> Path:
    """Generate a complete DOCX energy-analysis report.

    Parameters
    ----------
    data:
        A fully populated ``ReportData`` instance (see ``report.models``).

    Returns
    -------
    Path
        Absolute path to the generated ``.docx`` file (inside a temp dir).
    """
    doc = Document()

    # --- Base style ---------------------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # === SECTION 1: Cover page ==============================================
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AzzeroCO2 Energy")
    run.font.size = Pt(28)
    run.font.color.rgb = PRIMARY_BLUE
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Report di Analisi Energetica")
    run.font.size = Pt(18)
    run.font.color.rgb = MUTED_TEXT

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"\n{data.analysis.name}\n").font.size = Pt(16)
    info.add_run(f"\nAnno di riferimento: {data.analysis.year}\n").font.size = Pt(12)
    info.add_run(f"Sito: {data.analysis.site.name}").font.size = Pt(12)
    if data.analysis.site.city:
        info.add_run(f" \u2014 {data.analysis.site.city}").font.size = Pt(12)

    doc.add_paragraph()
    footer_text = doc.add_paragraph()
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_text.add_run(f"Generato il {data.generated_at[:10]}")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED_TEXT

    doc.add_page_break()

    # === SECTION 2: Site information ========================================
    _add_heading(doc, "1. Informazioni sito")

    site = data.analysis.site
    _add_table(doc, ["Campo", "Valore"], [
        ["Nome", site.name],
        ["Indirizzo", site.address or "\u2014"],
        ["Citt\u00e0", site.city or "\u2014"],
        ["Provincia", site.province or "\u2014"],
        ["Coordinate", f"{site.lat}, {site.lng}" if site.lat else "\u2014"],
        ["Anno di riferimento", str(data.analysis.year)],
        ["WACC", _fmt_pct(data.analysis.wacc) if data.analysis.wacc else "\u2014"],
    ])

    # === SECTION 3: Current situation (AS-IS) ===============================
    _add_heading(doc, "2. Situazione attuale")

    doc.add_paragraph("Consumi energetici annuali:", style="List Bullet")
    _add_table(
        doc,
        ["Vettore energetico", "Consumo annuo (MWh)"],
        [
            [d.end_use_label, f"{d.annual_consumption_mwh:,.1f}"]
            for d in data.demands
        ],
    )

    total_mwh = sum(d.annual_consumption_mwh for d in data.demands)
    p = doc.add_paragraph()
    run = p.add_run(f"Consumo totale: {total_mwh:,.1f} MWh/anno")
    run.font.bold = True

    # Resources table
    if data.resources:
        doc.add_paragraph()
        doc.add_paragraph("Risorse energetiche disponibili:", style="List Bullet")
        _add_table(
            doc,
            ["Risorsa", "Prezzo (\u20ac/MWh)", "Fattore CO\u2082 (tCO\u2082/MWh)"],
            [
                [
                    r.resource_label,
                    f"{r.buying_price:,.2f}" if r.buying_price else "\u2014",
                    f"{r.co2_factor:,.4f}" if r.co2_factor else "\u2014",
                ]
                for r in data.resources
            ],
        )

    # --- Generate charts (used in multiple sections) -----------------------
    charts = generate_all_charts(data)

    # Energy mix chart
    if data.demands:
        doc.add_paragraph()
        _add_centered_image(doc, str(charts["energy_mix"]), Inches(5))

    doc.add_page_break()

    # === SECTION 4: Proposed interventions ==================================
    _add_heading(doc, "3. Interventi proposti")

    scenario = data.scenario

    p = doc.add_paragraph()
    p.add_run("Scenario: ").font.bold = True
    p.add_run(scenario.name)

    p = doc.add_paragraph()
    p.add_run("Obiettivo: ").font.bold = True
    p.add_run(
        "Minimizzazione costi" if scenario.objective == "cost" else "Decarbonizzazione"
    )

    doc.add_paragraph()

    if scenario.tech_results:
        _add_table(
            doc,
            [
                "Tecnologia",
                "Categoria",
                "Capacit\u00e0 (kW)",
                "Produzione (MWh/a)",
                "CAPEX (\u20ac)",
                "Risparmio (\u20ac/a)",
            ],
            [
                [
                    t.technology_name,
                    t.category,
                    f"{t.optimal_capacity_kw:,.1f}",
                    f"{t.annual_production_mwh:,.1f}",
                    _fmt_eur(t.capex),
                    _fmt_eur(t.annual_savings),
                ]
                for t in scenario.tech_results
            ],
        )

        # CAPEX breakdown chart
        _add_centered_image(doc, str(charts["capex_breakdown"]), Inches(5.5))

    doc.add_page_break()

    # === SECTION 5: Financial analysis ======================================
    _add_heading(doc, "4. Analisi finanziaria")

    _add_table(doc, ["KPI", "Valore"], [
        ["CAPEX totale", _fmt_eur(scenario.total_capex)],
        ["OPEX annuo", _fmt_eur(scenario.total_opex_annual)],
        ["Risparmio annuo", _fmt_eur(scenario.total_savings_annual)],
        [
            "Payback",
            f"{scenario.payback_years:.1f} anni" if scenario.payback_years else "\u2014",
        ],
        ["IRR", _fmt_pct(scenario.irr)],
        ["NPV", _fmt_eur(scenario.npv) if scenario.npv else "\u2014"],
        ["Riduzione CO\u2082", _fmt_pct(scenario.co2_reduction_percent)],
    ])

    # CAPEX vs savings chart
    doc.add_paragraph()
    _add_centered_image(doc, str(charts["capex_vs_savings"]), Inches(5.5))

    # Cashflow chart
    doc.add_paragraph()
    _add_centered_image(doc, str(charts["cashflow"]), Inches(5.5))

    doc.add_page_break()

    # === SECTION 6: Conclusions =============================================
    _add_heading(doc, "5. Conclusioni e prossimi passi")

    doc.add_paragraph(
        f"L\u2019analisi energetica condotta per il sito \u00ab{data.analysis.site.name}\u00bb "
        f"ha evidenziato un consumo totale di {total_mwh:,.0f} MWh/anno.",
    )

    if scenario.co2_reduction_percent:
        doc.add_paragraph(
            f"Lo scenario \u00ab{scenario.name}\u00bb consente una riduzione delle emissioni "
            f"di CO\u2082 pari al {scenario.co2_reduction_percent * 100:.1f}%.",
        )

    doc.add_paragraph(
        f"L\u2019investimento complessivo ammonta a {_fmt_eur(scenario.total_capex)}, "
        f"con un risparmio annuo stimato di {_fmt_eur(scenario.total_savings_annual)}.",
    )

    if scenario.payback_years:
        doc.add_paragraph(
            f"Il tempo di rientro dell\u2019investimento (payback) \u00e8 stimato "
            f"in {scenario.payback_years:.1f} anni.",
        )

    doc.add_paragraph()
    doc.add_paragraph("Prossimi passi consigliati:", style="List Bullet")
    doc.add_paragraph(
        "Verifica dettagliata dei preventivi con i fornitori", style="List Bullet 2"
    )
    doc.add_paragraph(
        "Analisi degli incentivi e sussidi disponibili", style="List Bullet 2"
    )
    doc.add_paragraph(
        "Pianificazione delle fasi di implementazione", style="List Bullet 2"
    )
    doc.add_paragraph(
        "Monitoraggio dei risultati post-intervento", style="List Bullet 2"
    )

    # --- Branding footer ----------------------------------------------------
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "Report generato da AzzeroCO2 Energy \u2014 il clima nelle nostre mani"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED_TEXT
    run.font.italic = True

    # --- Save document ------------------------------------------------------
    safe_name = data.analysis.name.replace(" ", "_")
    output_path = (
        Path(tempfile.mkdtemp(prefix="azzeroco2_report_"))
        / f"report_{safe_name}.docx"
    )
    doc.save(str(output_path))

    # Cleanup chart temp files
    for chart_path in charts.values():
        chart_path.unlink(missing_ok=True)

    return output_path
